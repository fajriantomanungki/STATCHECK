import textwrap
import uuid
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.models.release import Release, ReleaseMinutes
from app.services.file_storage import storage_root


def build_minutes_text(release: Release, minutes: ReleaseMinutes) -> str:
    date_label = release.tanggal_rilis.strftime("%d-%m-%Y")
    time_label = release.waktu_rilis.strftime("%H:%M")
    lines = [
        "NOTULEN RILIS BERITA RESMI STATISTIK", "",
        f"Judul: {release.judul_rilis}", f"Tanggal: {date_label}",
        f"Waktu: {time_label}", f"Tempat: {release.tempat}", "",
        "DAFTAR BRS",
    ]
    lines.extend(f"{index}. {link.brs.nama_brs}" for index, link in enumerate(release.brs_links, 1))
    lines.extend(["", "PESERTA"])
    lines.extend(
        f"{index}. {guest.nama} - {guest.instansi}{f' ({guest.jabatan})' if guest.jabatan else ''}"
        for index, guest in enumerate(sorted(release.guests, key=lambda item: item.nama), 1)
    )
    if not release.guests:
        lines.append("Belum ada peserta terdaftar.")
    lines.extend(["", "PEMBUKAAN", minutes.opening or "-", "", "POKOK PEMBAHASAN", minutes.discussion or "-"])
    lines.extend(["", "SESI TANYA JAWAB"])
    finalized = [item for item in release.qna_items if item.final_answer]
    for index, item in enumerate(finalized, 1):
        asker = f"{item.guest.nama} - {item.guest.instansi}" if item.guest else "Penanya tidak tercatat"
        lines.extend([f"{index}. {asker}", f"Pertanyaan: {item.question}", f"Jawaban: {item.final_answer}", ""])
    if not finalized:
        lines.append("Tidak ada tanya jawab yang dicatat.")
    lines.extend(["", "CATATAN", minutes.notes or "-", "", "KESIMPULAN", minutes.conclusion or "-"])
    return "\n".join(lines)


def _write_docx(text: str, path: Path) -> None:
    document = DocxDocument()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for index, line in enumerate(text.splitlines()):
        paragraph = document.add_paragraph()
        if index == 0:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line)
            run.bold = True
            run.font.size = Pt(14)
        elif line in {"DAFTAR BRS", "PESERTA", "PEMBUKAAN", "POKOK PEMBAHASAN", "SESI TANYA JAWAB", "CATATAN", "KESIMPULAN"}:
            run = paragraph.add_run(line)
            run.bold = True
            run.font.size = Pt(11)
        else:
            paragraph.add_run(line)
    document.save(path)


def _pdf_safe(text: str) -> str:
    return text.replace("—", "-").replace("–", "-").replace("“", '"').replace("”", '"')


def _write_pdf(text: str, path: Path) -> None:
    document = fitz.open()
    page = document.new_page(width=595, height=842)
    y = 55
    headings = {"DAFTAR BRS", "PESERTA", "PEMBUKAAN", "POKOK PEMBAHASAN", "SESI TANYA JAWAB", "CATATAN", "KESIMPULAN"}
    for original in text.splitlines():
        line = _pdf_safe(original)
        wrapped = textwrap.wrap(line, width=92, break_long_words=False) or [""]
        is_heading = line in headings or line == "NOTULEN RILIS BERITA RESMI STATISTIK"
        for part in wrapped:
            if y > 795:
                page = document.new_page(width=595, height=842)
                y = 55
            page.insert_text((50, y), part, fontname="hebo" if is_heading else "helv", fontsize=11 if is_heading else 9)
            y += 16 if is_heading else 13
        if not line:
            y += 4
    document.save(path)
    document.close()


def generate_minutes_files(release: Release, minutes: ReleaseMinutes) -> tuple[str, str, str]:
    text = build_minutes_text(release, minutes)
    relative_dir = Path("minutes") / str(release.id)
    absolute_dir = storage_root() / relative_dir
    absolute_dir.mkdir(parents=True, exist_ok=True)
    suffix = uuid.uuid4().hex[:8]
    docx_relative = relative_dir / f"notulen_{suffix}.docx"
    pdf_relative = relative_dir / f"notulen_{suffix}.pdf"
    _write_docx(text, storage_root() / docx_relative)
    _write_pdf(text, storage_root() / pdf_relative)
    return text, docx_relative.as_posix(), pdf_relative.as_posix()

